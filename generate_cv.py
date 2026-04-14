from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_cv():
    doc = Document()

    # Title / Name
    name = doc.add_heading('Nguyễn Văn A', 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Basic Info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Email: test@example.com | SĐT: 0901 234 567 | Hà Nội, Việt Nam\nLinkedIn: linkedin.com/in/nguyenvana')
    run.font.size = Pt(10)

    # Summary
    doc.add_heading('Tóm tắt chuyên môn', level=1)
    doc.add_paragraph(
        'Kỹ sư phần mềm với hơn 3 năm kinh nghiệm phát triển các ứng dụng web quy mô lớn. '
        'Chuyên gia về Fullstack Development (React/Node.js) và xử lý dữ liệu với Python. '
        'Luôn hướng tới việc tạo ra mã nguồn sạch, tối ưu và giao diện người dùng tinh tế.'
    )

    # Skills
    doc.add_heading('Kỹ năng chính', level=1)
    doc.add_paragraph('Lập trình: JavaScript (ES6+), Python, HTML5, CSS3/Tailwind', style='List Bullet')
    doc.add_paragraph('Frameworks: React.js, Express.js, FastAPI', style='List Bullet')
    doc.add_paragraph('Database: MySQL, SQLite, MongoDB', style='List Bullet')
    doc.add_paragraph('Công cụ: Git, Docker, AWS', style='List Bullet')

    # Experience
    doc.add_heading('Kinh nghiệm làm việc', level=1)
    
    # Job 1
    p = doc.add_paragraph()
    run = p.add_run('Công ty Công nghệ ABC - Software Engineer')
    run.bold = True
    p.add_run('\n01/2022 - Hiện tại')
    doc.add_paragraph(
        '• Xây dựng và bảo trì hệ thống quản lý tuyển dụng thông minh.\n'
        '• Tối ưu hóa hiệu suất frontend, giảm thời gian tải trang xuống 40%.\n'
        '• Tích hợp AI (Gemini/OpenAI) để phân tích CV tự động.',
        style='Normal'
    )

    # Education
    doc.add_heading('Học vấn', level=1)
    p = doc.add_paragraph()
    run = p.add_run('Đại học Bách Khoa Hà Nội - Kỹ thuật Phần mềm')
    run.bold = True
    p.add_run('\nGPA: 3.6/4.0 | 2018 - 2022')

    # Projects
    doc.add_heading('Dự án tiêu biểu', level=1)
    p = doc.add_paragraph()
    run = p.add_run('AI CV Analyst (Dự án cá nhân)')
    run.bold = True
    doc.add_paragraph(
        'Hệ thống phân tích CV tự động sử dụng React, Node.js và Python FastAPI. '
        'Dự án đạt hơn 100 stars trên GitHub.'
    )

    # Save
    file_path = "CV_NguyenVanA.docx"
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    path = create_cv()
    print(f"CV created successfully at: {os.path.abspath(path)}")
